#!/usr/bin/env python3
"""
Enhanced Document Loader with Metadata Extraction
Handles all document types for litigation intelligence
"""

from pathlib import Path
from dotenv import load_dotenv
import os

env_path = Path(__file__).parent / '.env'
load_dotenv(dotenv_path=env_path, override=True)

from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import hashlib
import json

# PDF libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

import PyPDF2
import docx
import chardet
import re


class DocumentLoader:
    """Advanced document loading with metadata extraction"""
    
    def __init__(self, config):
        self.config = config
        self.load_stats = {
            'total_loaded': 0,
            'failed_loads': 0,
            'by_type': {}
        }
    
    def load_directory(self, 
                      directory: Path,
                      max_docs: Optional[int] = None,
                      doc_types: List[str] = None) -> List[Dict]:
        """
        Load all documents from directory
        
        Args:
            directory: Path to document directory
            max_docs: Maximum documents to load
            doc_types: List of extensions to load (e.g., ['.pdf', '.docx'])
        
        Returns:
            List of document dictionaries with content and metadata
        """
        
        if not directory.exists():
            print(f"Warning: Directory {directory} does not exist")
            return []
        
        # Default document types
        if doc_types is None:
            doc_types = ['.pdf', '.docx', '.doc', '.txt', '.json', '.html', '.md']
        
        documents = []
        files_processed = 0
        
        # Get all matching files
        all_files = []
        for doc_type in doc_types:
            all_files.extend(directory.glob(f"**/*{doc_type}"))
        
        # Sort for consistent ordering
        all_files.sort()
        
        # Limit if specified
        if max_docs:
            all_files = all_files[:max_docs]
        
        print(f"Loading {len(all_files)} documents from {directory}")
        
        for file_path in all_files:
            files_processed += 1
            
            if files_processed % 10 == 0:
                print(f"  Progress: {files_processed}/{len(all_files)}")
            
            try:
                document = self.load_document(file_path)
                if document:
                    documents.append(document)
                    self.load_stats['total_loaded'] += 1
            except Exception as e:
                print(f"  Failed to load {file_path.name}: {e}")
                self.load_stats['failed_loads'] += 1
        
        print(f"✅ Loaded {len(documents)} documents successfully")
        return documents
    
    def load_document(self, file_path: Path) -> Optional[Dict]:
        """
        Load single document with metadata extraction
        
        Returns:
            Document dictionary with content and metadata
        """
        
        if not file_path.exists():
            return None
        
        # Generate document ID
        doc_id = self._generate_doc_id(file_path)
        
        # Load based on file type
        extension = file_path.suffix.lower()
        
        content = None
        extraction_metadata = {}
        
        if extension == '.pdf':
            content, extraction_metadata = self._load_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            content = self._load_docx(file_path)
        elif extension == '.json':
            content = self._load_json(file_path)
        elif extension in ['.txt', '.md', '.html']:
            content = self._load_text(file_path)
        else:
            content = self._load_text(file_path)  # Try as text
        
        if not content:
            return None
        
        # Extract metadata
        metadata = self._extract_metadata(file_path, content, extraction_metadata)
        
        # Create document entry
        document = {
            'id': doc_id,
            'filename': file_path.name,
            'path': str(file_path),
            'content': content,
            'metadata': metadata
        }
        
        # Track statistics
        if extension not in self.load_stats['by_type']:
            self.load_stats['by_type'][extension] = 0
        self.load_stats['by_type'][extension] += 1
        
        return document
    
    def _load_pdf(self, file_path: Path) -> Tuple[Optional[str], Dict]:
        """Load PDF with best available method"""
        
        metadata = {
            'extraction_method': None,
            'page_count': 0,
            'has_tables': False,
            'has_images': False,
            'extraction_quality': 'unknown'
        }
        
        # Try PyMuPDF first (best quality)
        if PYMUPDF_AVAILABLE:
            try:
                import fitz
                doc = fitz.open(str(file_path))
                text = ""
                
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    text += f"\n[PAGE {page_num + 1}]\n{page_text}"
                    
                    # Check for tables and images
                    if page.get_tables():
                        metadata['has_tables'] = True
                    if page.get_images():
                        metadata['has_images'] = True
                
                metadata['page_count'] = len(doc)
                metadata['extraction_method'] = 'PyMuPDF'
                metadata['extraction_quality'] = 'high'
                doc.close()
                
                if text.strip():
                    return text, metadata
                    
            except Exception as e:
                pass
        
        # Try pdfplumber (good for tables)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        text += f"\n[PAGE {i + 1}]\n{page_text}"
                        
                        if page.extract_tables():
                            metadata['has_tables'] = True
                    
                    metadata['page_count'] = len(pdf.pages)
                    metadata['extraction_method'] = 'pdfplumber'
                    metadata['extraction_quality'] = 'medium'
                    
                    if text.strip():
                        return text, metadata
                        
            except Exception:
                pass
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += f"\n[PAGE {page_num + 1}]\n{page_text}"
                
                metadata['page_count'] = len(pdf_reader.pages)
                metadata['extraction_method'] = 'PyPDF2'
                metadata['extraction_quality'] = 'basic'
                
                return text, metadata
                
        except Exception as e:
            print(f"All PDF extraction failed for {file_path}: {e}")
            return None, metadata
    
    def _load_docx(self, file_path: Path) -> Optional[str]:
        """Load Word document"""
        
        try:
            doc = docx.Document(str(file_path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return '\n'.join(paragraphs)
            
        except Exception as e:
            print(f"Failed to load DOCX {file_path}: {e}")
            return None
    
    def _load_json(self, file_path: Path) -> Optional[str]:
        """Load JSON file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Convert to formatted string
                return json.dumps(data, indent=2)
        except Exception as e:
            print(f"Failed to load JSON {file_path}: {e}")
            return None
    
    def _load_text(self, file_path: Path) -> Optional[str]:
        """Load text file with encoding detection"""
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # Load with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            print(f"Failed to load text {file_path}: {e}")
            return None
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID"""
        
        # Use file path hash for consistency
        path_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8].upper()
        
        # Prefix based on location
        if 'legal' in str(file_path).lower():
            prefix = 'LEGAL'
        elif 'case' in str(file_path).lower():
            prefix = 'CASE'
        elif 'disclosure' in str(file_path).lower():
            prefix = 'DISC'
        else:
            prefix = 'DOC'
        
        return f"{prefix}_{path_hash}"
    
    def _extract_metadata(self, 
                         file_path: Path,
                         content: str,
                         extraction_metadata: Dict) -> Dict:
        """Extract comprehensive metadata from document"""
        
        metadata = {
            'file_size_mb': file_path.stat().st_size / (1024 * 1024),
            'modified_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
            'extension': file_path.suffix.lower(),
            'char_count': len(content),
            'word_count': len(content.split()),
            'line_count': content.count('\n')
        }
        
        # Add extraction metadata
        metadata.update(extraction_metadata)
        
        # Extract dates
        dates = self._extract_dates(content)
        metadata['dates_found'] = dates[:10]  # Top 10 dates
        metadata['has_dates'] = len(dates) > 0
        
        # Extract amounts
        amounts = self._extract_amounts(content)
        metadata['amounts_found'] = amounts[:10]  # Top 10 amounts
        metadata['has_amounts'] = len(amounts) > 0
        
        # Extract entities
        entities = self._extract_entities(content)
        metadata['entities'] = entities
        
        # Document classification
        metadata['classification'] = self._classify_document(file_path.name, content)
        
        # Content flags
        metadata['has_tables'] = bool(re.search(r'\t|\|', content))
        metadata['has_emails'] = bool(re.findall(r'\b[\w.-]+@[\w.-]+\.\w+\b', content))
        
        return metadata
    
    def _extract_dates(self, content: str) -> List[str]:
        """Extract dates from content"""
        
        dates = []
        
        # Multiple date patterns
        patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',  # DD/MM/YYYY or MM/DD/YYYY
            r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',    # YYYY-MM-DD
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b',  # Month DD, YYYY
            r'\b\d{1,2} (?:January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            dates.extend(matches)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_dates = []
        for date in dates:
            if date not in seen:
                seen.add(date)
                unique_dates.append(date)
        
        return unique_dates
    
    def _extract_amounts(self, content: str) -> List[str]:
        """Extract monetary amounts from content"""
        
        amounts = []
        
        # Amount patterns
        patterns = [
            r'[£$€]\s*[\d,]+(?:\.\d{2})?(?:\s*(?:million|billion|k|m|bn))?',
            r'[\d,]+(?:\.\d{2})?\s*(?:GBP|USD|EUR)',
            r'[\d,]+(?:\.\d{2})?\s*(?:pounds?|dollars?|euros?)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            amounts.extend(matches)
        
        # Clean and deduplicate
        unique_amounts = list(set(amounts))
        return unique_amounts[:20]  # Limit to 20
    
    def _extract_entities(self, content: str) -> Dict[str, List[str]]:
        """Extract named entities from content"""
        
        entities = {
            'people': [],
            'companies': [],
            'emails': []
        }
        
        # Extract people (simple pattern)
        people_pattern = r'\b([A-Z][a-z]+ (?:[A-Z]\. )?[A-Z][a-z]+)\b'
        people = re.findall(people_pattern, content)
        entities['people'] = list(set(people))[:20]
        
        # Extract companies
        company_patterns = [
            r'\b\w+\s+(?:Ltd|Limited|Inc|LLC|LLP|Plc|Corp|Corporation)\b',
            r'\b\w+\s+(?:Capital|Holdings|Partners|Group|Ventures)\b'
        ]
        
        companies = []
        for pattern in company_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            companies.extend(matches)
        entities['companies'] = list(set(companies))[:20]
        
        # Extract emails
        email_pattern = r'\b[\w.-]+@[\w.-]+\.\w+\b'
        emails = re.findall(email_pattern, content)
        entities['emails'] = list(set(emails))[:10]
        
        return entities
    
    def _classify_document(self, filename: str, content: str) -> str:
        """Classify document type based on content and filename"""
        
        filename_lower = filename.lower()
        content_lower = content[:5000].lower()  # Check first 5000 chars
        
        # Check filename patterns
        if 'contract' in filename_lower or 'agreement' in filename_lower:
            return 'contract'
        elif 'email' in filename_lower or 'correspondence' in filename_lower:
            return 'correspondence'
        elif 'invoice' in filename_lower or 'receipt' in filename_lower:
            return 'financial'
        elif 'statement' in filename_lower or 'witness' in filename_lower:
            return 'witness_statement'
        elif 'minutes' in filename_lower or 'meeting' in filename_lower:
            return 'minutes'
        
        # Check content patterns
        if 'whereas' in content_lower and 'agreement' in content_lower:
            return 'contract'
        elif 'from:' in content_lower and 'to:' in content_lower:
            return 'email'
        elif 'invoice' in content_lower or 'payment' in content_lower:
            return 'financial'
        elif 'i state' in content_lower or 'witness statement' in content_lower:
            return 'witness_statement'
        elif 'minutes of' in content_lower:
            return 'minutes'
        
        return 'general'
    
    def get_statistics(self) -> Dict:
        """Get loading statistics"""
        
        return {
            'total_loaded': self.load_stats['total_loaded'],
            'failed_loads': self.load_stats['failed_loads'],
            'by_type': self.load_stats['by_type'],
            'success_rate': (
                self.load_stats['total_loaded'] / 
                max(1, self.load_stats['total_loaded'] + self.load_stats['failed_loads'])
            )
        }