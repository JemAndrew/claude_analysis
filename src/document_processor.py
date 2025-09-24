# document_processor.py - FIXED VERSION
"""
Core document processing engine that handles phase-by-phase analysis.
Now handles PDFs, DOCX, TXT, and other file types properly.
"""

import json
from typing import List, Dict, Optional
from datetime import datetime
from pathlib import Path
import PyPDF2
import docx
import chardet

from api_client import ClaudeAPIClient
from master_prompt import get_master_prompt
from phase_prompts import (
    get_phase_prompt, 
    update_learning_prompt, 
    get_learning_generator_prompt,
    should_generate_learning,
    get_phase_description,
    get_all_phases
)
from knowledge_manage import KnowledgeManager

class DocumentProcessor:
    """Processes documents through all analysis phases"""
    
    def __init__(self):
        """Initialise the document processor"""
        self.api_client = ClaudeAPIClient()
        self.knowledge_manage = KnowledgeManager()
        self.current_phase = None
        self.documents = []
        
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from any file type
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            # PDF files
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            
            # Word documents
            elif extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            
            # Text-based files
            elif extension in ['.txt', '.md', '.rtf', '.csv', '.log']:
                return self._extract_from_text(file_path)
            
            # JSON files
            elif extension == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            
            # HTML/XML files
            elif extension in ['.html', '.htm', '.xml']:
                return self._extract_from_text(file_path)
            
            # Default: try to read as text
            else:
                print(f"  ⚠️ Unknown file type {extension}, attempting text extraction...")
                return self._extract_from_text(file_path)
                
        except Exception as e:
            print(f"  ❌ Error extracting from {path.name}: {e}")
            return f"[Error reading file: {path.name}]"
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                print(f"    Extracting {num_pages} pages from PDF...")
                
                for page_num in range(num_pages):
                    try:
                        page = reader.pages[page_num]
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        print(f"      Warning: Could not extract page {page_num + 1}: {e}")
                
                if not text_parts:
                    print(f"      Warning: No text extracted from PDF, might be scanned/image-based")
                    return "[PDF appears to be image-based - OCR required]"
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            print(f"    Error reading PDF: {e}")
            # Try alternative method with PyMuPDF if available
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except ImportError:
                return f"[Error reading PDF: {e}]"
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n\n".join(full_text)
            
        except Exception as e:
            print(f"    Error reading DOCX: {e}")
            return f"[Error reading DOCX: {e}]"
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from text-based files with encoding detection"""
        try:
            # First, try UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                pass
            
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                
                if encoding:
                    return raw_data.decode(encoding)
                else:
                    # Try common encodings
                    for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            return raw_data.decode(enc)
                        except:
                            continue
                    
                    return "[Could not decode file - unknown encoding]"
                    
        except Exception as e:
            print(f"    Error reading text file: {e}")
            return f"[Error reading file: {e}]"
    
    def load_documents(self, document_paths: List[str]) -> bool:
        """
        Load documents from file paths - handles all file types
        
        Args:
            document_paths: List of paths to document files
            
        Returns:
            Success status
        """
        try:
            self.documents = []
            successful = 0
            failed = 0
            
            for path in document_paths:
                path_obj = Path(path)
                
                if not path_obj.exists():
                    print(f"  ⚠️ File not found: {path}")
                    failed += 1
                    continue
                
                print(f"  Loading: {path_obj.name}")
                
                # Extract text based on file type
                content = self.extract_text_from_file(str(path_obj))
                
                if content and not content.startswith("[Error"):
                    self.documents.append({
                        'path': str(path_obj),
                        'content': content,
                        'filename': path_obj.name
                    })
                    successful += 1
                else:
                    print(f"    ❌ Failed to extract content from {path_obj.name}")
                    failed += 1
            
            print(f"\n📊 Loading Summary:")
            print(f"  ✅ Successfully loaded: {successful} documents")
            if failed > 0:
                print(f"  ❌ Failed to load: {failed} documents")
            
            return successful > 0
            
        except Exception as e:
            print(f"Error loading documents: {e}")
            return False
    
    def process_phase(self, phase_num: str, documents: Optional[List[Dict]] = None) -> Dict:
        """
        Process documents for a specific phase
        
        Args:
            phase_num: Phase identifier (0A, 0B, 1-7)
            documents: Optional document list (uses self.documents if not provided)
            
        Returns:
            Analysis results dictionary
        """
        print(f"\n{'='*60}")
        print(f"PROCESSING PHASE {phase_num}: {get_phase_description(phase_num)}")
        print(f"{'='*60}")
        
        self.current_phase = phase_num
        docs_to_process = documents or self.documents
        
        if not docs_to_process:
            print("ERROR: No documents to process!")
            return {}
        
        # Get previous knowledge for context
        previous_knowledge = self.knowledge_manage.get_previous_knowledge(phase_num)
        
        # 1. Get the master forensic overlay (always applied)
        master = get_master_prompt()
        
        # 2. Get phase-specific prompt (includes any learned components if re-running)
        phase_prompt = get_phase_prompt(phase_num)
        
        # 3. Combine prompts with previous knowledge
        full_prompt = self._build_full_prompt(master, phase_prompt, previous_knowledge)
        
        # 4. Prepare documents for analysis
        document_texts = []
        for doc in docs_to_process:
            text = doc['content']
            # Truncate very long documents to avoid token limits
            if len(text) > 50000:
                text = text[:45000] + "\n\n[...document truncated for length...]\n\n" + text[-5000:]
            document_texts.append(text)
        
        print(f"Sending {len(document_texts)} documents for analysis...")
        
        # 5. Send to Claude for initial analysis
        # Fix the phase format for api_client
        api_phase = f"phase_{phase_num.lower()}" if phase_num else None
        
        response = self.api_client.analyse_documents(
            documents=document_texts,
            prompt=full_prompt,
            phase=api_phase,
            context=previous_knowledge
        )
        
        if not response:
            print(f"Error: No response received for phase {phase_num}")
            return {}
        
        # 6. For phases 1-6, generate and apply Claude's learning enhancement
        if should_generate_learning(phase_num):
            print(f"Generating autonomous enhancement for phase {phase_num}...")
            enhanced_response = self._apply_learning_enhancement(
                phase_num, 
                response, 
                docs_to_process, 
                previous_knowledge
            )
            
            # Combine both analyses
            final_response = {
                'initial_analysis': response,
                'enhanced_analysis': enhanced_response,
                'combined_insights': f"{response}\n\nENHANCED ANALYSIS:\n{enhanced_response}"
            }
        else:
            final_response = {
                'analysis': response
            }
        
        # 7. Store the knowledge
        phase_result = {
            'phase': phase_num,
            'description': get_phase_description(phase_num),
            'timestamp': datetime.now().isoformat(),
            'documents_analysed': len(docs_to_process),
            'findings': final_response
        }
        
        self.knowledge_manage.store_phase_knowledge(phase_num, phase_result)
        
        print(f"Phase {phase_num} complete. Knowledge stored.")
        
        return phase_result
    
    def _build_full_prompt(self, master: str, phase_prompt: str, previous_knowledge: Dict) -> str:
        """
        Build the complete prompt for analysis
        
        Args:
            master: Master forensic prompt
            phase_prompt: Phase-specific prompt
            previous_knowledge: Knowledge from previous phases
            
        Returns:
            Combined prompt string
        """
        prompt_parts = [
            master,
            "\nCURRENT PHASE FOCUS:",
            phase_prompt
        ]
        
        if previous_knowledge:
            prompt_parts.extend([
                "\nPREVIOUS PHASE KNOWLEDGE:",
                "You have access to the following insights from previous phases:"
            ])
            
            for phase, knowledge in previous_knowledge.items():
                phase_desc = get_phase_description(phase)
                prompt_parts.append(f"\n[{phase} - {phase_desc}]:")
                
                # Summarise key findings to avoid token overflow
                if isinstance(knowledge, dict) and 'findings' in knowledge:
                    findings = knowledge['findings']
                    if isinstance(findings, dict):
                        # Handle combined insights from enhanced phases
                        if 'combined_insights' in findings:
                            prompt_parts.append(findings['combined_insights'][:3000])
                        elif 'analysis' in findings:
                            prompt_parts.append(findings['analysis'][:3000])
                    else:
                        prompt_parts.append(str(findings)[:3000])
        
        prompt_parts.append("\nAnalyse the provided documents through both the master forensic lens "
                          "and the specific phase focus. Find what others would miss. "
                          "Be aggressive in identifying anything that damages Process Holdings' case.")
        
        return "\n".join(prompt_parts)
    
    def _apply_learning_enhancement(
        self, 
        phase_num: str, 
        initial_response: str, 
        documents: List[Dict], 
        previous_knowledge: Dict
    ) -> str:
        """
        Generate and apply Claude's self-generated enhancement
        
        Args:
            phase_num: Current phase
            initial_response: Initial analysis response
            documents: Documents being analysed
            previous_knowledge: Previous phase knowledge
            
        Returns:
            Enhanced analysis response
        """
        # Ask Claude to generate enhancement based on initial findings
        enhancement_prompt = f"""
        You just completed Phase {phase_num} analysis with these findings:
        {initial_response[:3000]}
        
        {get_learning_generator_prompt()}
        
        Be specific about:
        - Document references that need deeper investigation
        - Patterns you've detected that could be significant
        - Connections to previous phase findings
        - Areas where Process Holdings seems vulnerable
        """
        
        # Fix phase format for API
        api_phase = f"phase_{phase_num.lower()}_learning"
        
        # Get Claude's self-generated enhancement
        learning = self.api_client.analyse_documents(
            documents=[],  # No documents needed for prompt generation
            prompt=enhancement_prompt,
            phase=api_phase
        )
        
        # Store the learning for future use
        update_learning_prompt(phase_num, learning)
        
        # Re-run the phase with enhancement
        master = get_master_prompt()
        enhanced_phase_prompt = get_phase_prompt(phase_num, include_learning=True)
        
        enhanced_full_prompt = f"""
        {master}
        
        ENHANCED PHASE FOCUS:
        {enhanced_phase_prompt}
        
        PREVIOUS PHASE KNOWLEDGE:
        {json.dumps(previous_knowledge, indent=2)[:5000]}
        
        This is your SECOND PASS with your own insights. Go deeper. Be more aggressive.
        Find the evidence that destroys Process Holdings' case.
        """
        
        # Get document texts
        document_texts = [doc['content'][:50000] for doc in documents]
        
        # Fix phase format
        api_phase_enhanced = f"phase_{phase_num.lower()}_enhanced"
        
        # Second, enhanced analysis
        enhanced_response = self.api_client.analyse_documents(
            documents=document_texts,
            prompt=enhanced_full_prompt,
            phase=api_phase_enhanced,
            context=previous_knowledge
        )
        
        return enhanced_response
    
    # Keep all your other methods unchanged...
    def run_full_analysis(self) -> Dict:
        """
        Run all phases in sequence
        
        Returns:
            Complete analysis results
        """
        print("\n" + "="*60)
        print("STARTING FULL ANALYSIS FOR LISMORE CAPITAL")
        print("="*60)
        
        all_phases = get_all_phases()
        results = {}
        
        for phase in all_phases:
            phase_result = self.process_phase(phase)
            results[phase] = phase_result
            
            # Brief pause between phases (optional)
            print(f"\nPhase {phase} complete. Moving to next phase...")
        
        print("\n" + "="*60)
        print("FULL ANALYSIS COMPLETE")
        print("="*60)
        
        # Generate summary
        summary = self.knowledge_manage.generate_summary()
        results['summary'] = summary
        
        return results
    
    def run_single_phase(self, phase_num: str) -> Dict:
        """
        Run a single phase
        
        Args:
            phase_num: Phase to run
            
        Returns:
            Phase results
        """
        return self.process_phase(phase_num)
    
    def get_status(self) -> Dict:
        """Get current processing status"""
        return {
            'current_phase': self.current_phase,
            'documents_loaded': len(self.documents),
            'phases_completed': self.knowledge_manage.get_completed_phases(),
            'total_phases': len(get_all_phases())
        }